from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DacTa_UseCase_CarShop.docx"


use_cases = [
    {
        "id": "UC01",
        "name": "Dang ky tai khoan",
        "actor": "Nguoi dung (Khach hang)",
        "description": "Use Case nay cho phep nguoi dung nhap thong tin ca nhan can thiet de tao mot tai khoan moi va dang ky su dung he thong.",
        "pre": ["Nguoi dung chua co tai khoan trong he thong."],
        "post": [
            "Tai khoan moi duoc tao trong he thong.",
            "Thong tin nguoi dung duoc luu vao co so du lieu.",
        ],
        "main": [
            "Nguoi dung truy cap trang dang ky.",
            "He thong hien thi form dang ky tai khoan.",
            "Nguoi dung nhap cac thong tin can thiet: ho ten, email, mat khau, xac nhan mat khau, so dien thoai.",
            "Nguoi dung nhan nut Dang ky.",
            "He thong kiem tra tinh hop le cua thong tin da nhap.",
            "Neu thong tin hop le, he thong tao tai khoan moi trong co so du lieu.",
            "He thong thong bao dang ky thanh cong.",
            "He thong chuyen nguoi dung den trang dang nhap.",
        ],
        "alt": [
            ("A1: Email da ton tai", [
                "He thong kiem tra thong tin dang ky.",
                "Phat hien email da duoc su dung.",
                "He thong hien thi thong bao loi.",
                "Nguoi dung nhap lai email khac.",
            ]),
            ("A2: Chua nhap day du thong tin", [
                "Nguoi dung nhan Dang ky nhung de trong mot hoac nhieu truong bat buoc.",
                "He thong hien thi thong bao yeu cau nhap day du thong tin.",
            ]),
            ("A3: Mat khau va xac nhan mat khau khong khop", [
                "He thong kiem tra thong tin mat khau.",
                "Phat hien mat khau va xac nhan mat khau khong giong nhau.",
                "He thong hien thi thong bao Mat khau xac nhan khong khop.",
            ]),
        ],
    },
    {
        "id": "UC02",
        "name": "Dang nhap he thong",
        "actor": "Khach hang, Nhan vien, Quan tri vien",
        "description": "Cho phep nguoi dung dang nhap bang email va mat khau de su dung cac chuc nang theo quyen.",
        "pre": ["Nguoi dung da co tai khoan hop le.", "Tai khoan dang o trang thai hoat dong."],
        "post": ["Nguoi dung nhan duoc phien dang nhap/JWT.", "He thong hien thi giao dien phu hop voi vai tro."],
        "main": [
            "Nguoi dung truy cap trang dang nhap.",
            "He thong hien thi form dang nhap.",
            "Nguoi dung nhap email va mat khau.",
            "Nguoi dung nhan nut Dang nhap.",
            "He thong xac thuc thong tin dang nhap.",
            "Neu hop le, he thong luu token dang nhap.",
            "He thong chuyen nguoi dung den trang chu hoac trang quan tri theo vai tro.",
        ],
        "alt": [
            ("A1: Sai email hoac mat khau", ["He thong tu choi dang nhap va hien thi thong bao loi."]),
            ("A2: Tai khoan bi khoa", ["He thong khong cho dang nhap va thong bao tai khoan khong kha dung."]),
        ],
    },
    {
        "id": "UC03",
        "name": "Xem va tim kiem o to",
        "actor": "Khach hang, Khach vang lai",
        "description": "Cho phep nguoi dung xem danh sach o to, loc theo hang xe, trang thai, khoang gia va tim kiem theo tu khoa.",
        "pre": ["He thong co du lieu o to dang ban."],
        "post": ["Nguoi dung xem duoc danh sach hoac chi tiet o to phu hop nhu cau."],
        "main": [
            "Nguoi dung truy cap trang danh sach o to.",
            "He thong hien thi danh sach o to co phan trang.",
            "Nguoi dung nhap tu khoa hoac chon bo loc.",
            "He thong gui yeu cau tim kiem/loc den backend.",
            "He thong hien thi ket qua phu hop.",
            "Nguoi dung chon mot o to de xem chi tiet.",
            "He thong hien thi thong tin chi tiet, hinh anh, tinh trang va thong tin kho neu co.",
        ],
        "alt": [
            ("A1: Khong co ket qua", ["He thong hien thi danh sach rong hoac thong bao khong tim thay san pham phu hop."]),
            ("A2: Loi tai du lieu", ["He thong hien thi thong bao loi va cho phep nguoi dung thu lai."]),
        ],
    },
    {
        "id": "UC04",
        "name": "Xem va tim kiem phu kien",
        "actor": "Khach hang, Khach vang lai",
        "description": "Cho phep nguoi dung xem danh sach phu kien, loc theo loai va tim kiem phu kien.",
        "pre": ["He thong co du lieu phu kien."],
        "post": ["Nguoi dung xem duoc danh sach/chi tiet phu kien."],
        "main": [
            "Nguoi dung truy cap trang phu kien.",
            "He thong hien thi danh sach phu kien.",
            "Nguoi dung chon loai phu kien hoac nhap tu khoa.",
            "He thong hien thi ket qua phu hop.",
            "Nguoi dung chon phu kien de xem chi tiet.",
        ],
        "alt": [
            ("A1: Phu kien het hang", ["He thong van hien thi thong tin nhung khong cho dat qua so luong ton kho."]),
            ("A2: Khong tim thay", ["He thong hien thi thong bao khong co phu kien phu hop."]),
        ],
    },
    {
        "id": "UC05",
        "name": "Xem dich vu va dat lich dich vu",
        "actor": "Khach hang",
        "description": "Cho phep khach hang xem danh sach dich vu, xem chi tiet va dat lich su dung dich vu tai chi nhanh.",
        "pre": ["Khach hang da dang nhap khi thuc hien dat lich.", "Dich vu va chi nhanh con hoat dong."],
        "post": ["Lich hen dich vu duoc tao trong he thong."],
        "main": [
            "Khach hang truy cap trang dich vu.",
            "He thong hien thi danh sach dich vu.",
            "Khach hang chon dich vu can xem.",
            "He thong hien thi chi tiet dich vu.",
            "Khach hang chon chi nhanh, ngay hen, gio hen va thong tin lien he.",
            "Khach hang xac nhan dat lich.",
            "He thong luu lich hen va thong bao thanh cong.",
        ],
        "alt": [
            ("A1: Chua dang nhap", ["He thong yeu cau khach hang dang nhap truoc khi dat lich."]),
            ("A2: Chua chon chi nhanh", ["He thong hien thi thong bao yeu cau chon chi nhanh."]),
            ("A3: Thieu thong tin lich hen", ["He thong yeu cau nhap day du ngay, gio va thong tin bat buoc."]),
        ],
    },
    {
        "id": "UC06",
        "name": "Dang ky lai thu o to",
        "actor": "Khach hang",
        "description": "Cho phep khach hang dat lich lai thu mot mau o to tai chi nhanh da chon.",
        "pre": ["Khach hang da dang nhap.", "O to co thong tin chi nhanh/kho phu hop."],
        "post": ["Lich lai thu duoc tao va cho nhan vien xu ly."],
        "main": [
            "Khach hang mo trang chi tiet o to.",
            "Khach hang chon chi nhanh muon lai thu.",
            "Khach hang mo form dang ky lai thu.",
            "He thong tu dien thong tin khach hang neu da dang nhap.",
            "Khach hang nhap ngay hen, gio hen va ghi chu neu co.",
            "Khach hang xac nhan dat lich.",
            "He thong tao lich hen loai LAI_THU va thong bao thanh cong.",
        ],
        "alt": [
            ("A1: Chua dang nhap", ["He thong chuyen khach hang den trang dang nhap."]),
            ("A2: Khong chon chi nhanh", ["He thong yeu cau chon chi nhanh truoc khi dat lich."]),
        ],
    },
    {
        "id": "UC07",
        "name": "Quan ly gio hang",
        "actor": "Khach hang",
        "description": "Cho phep khach hang them san pham vao gio, cap nhat so luong, chon kho, xoa san pham va xoa toan bo gio hang.",
        "pre": ["Khach hang da dang nhap hoac co phien gio hang hop le.", "San pham ton tai trong he thong."],
        "post": ["Gio hang duoc cap nhat theo thao tac cua khach hang."],
        "main": [
            "Khach hang chon san pham o to, phu kien hoac dich vu.",
            "Khach hang nhan Them vao gio hang.",
            "He thong them san pham vao gio hang.",
            "Khach hang mo trang gio hang.",
            "Khach hang thay doi so luong, chon kho hoac xoa mat hang.",
            "He thong cap nhat tong tien va trang thai gio hang.",
        ],
        "alt": [
            ("A1: So luong khong hop le", ["He thong tu choi cap nhat va hien thi thong bao loi."]),
            ("A2: San pham khong du ton kho", ["He thong thong bao so luong khong kha dung."]),
            ("A3: Gio hang rong", ["He thong hien thi trang gio hang rong va goi y tiep tuc mua sam."]),
        ],
    },
    {
        "id": "UC08",
        "name": "Dat hang va thanh toan",
        "actor": "Khach hang",
        "description": "Cho phep khach hang tao don hang tu gio hang, chon dia chi giao hang, tinh phi van chuyen GHN va thanh toan COD hoac VNPay.",
        "pre": ["Khach hang da dang nhap.", "Gio hang co san pham.", "Khach hang co dia chi giao hang hoac nhap dia chi moi."],
        "post": ["Don hang duoc tao trong he thong.", "Gio hang duoc xoa sau khi dat hang thanh cong.", "Thong tin thanh toan duoc ghi nhan neu co."],
        "main": [
            "Khach hang truy cap trang thanh toan.",
            "He thong tai gio hang va so dia chi cua khach hang.",
            "Khach hang chon dia chi co san hoac nhap dia chi giao hang moi.",
            "He thong goi GHN de tinh phi van chuyen.",
            "Khach hang chon phuong thuc thanh toan COD hoac VNPay.",
            "Khach hang xac nhan dat hang.",
            "He thong tao don hang va chi tiet don hang.",
            "Neu chon VNPay, he thong tao URL thanh toan va chuyen khach hang sang cong VNPay.",
            "He thong thong bao dat hang thanh cong va chuyen den lich su don hang.",
        ],
        "alt": [
            ("A1: Chua co dia chi giao hang", ["He thong yeu cau khach hang nhap day du dia chi moi."]),
            ("A2: Loi tinh phi van chuyen", ["He thong su dung phi mac dinh hoac hien thi thong bao loi tuy cau hinh."]),
            ("A3: Loi tao thanh toan VNPay", ["He thong thong bao loi VNPay; don hang co the van duoc ghi nhan."]),
            ("A4: Gio hang rong", ["He thong chuyen khach hang ve trang gio hang."]),
        ],
    },
    {
        "id": "UC09",
        "name": "Xem lich su va chi tiet don hang",
        "actor": "Khach hang",
        "description": "Cho phep khach hang xem cac don hang da dat va theo doi trang thai xu ly/giao hang.",
        "pre": ["Khach hang da dang nhap.", "Khach hang co don hang trong he thong."],
        "post": ["Khach hang nam duoc trang thai don hang va thong tin thanh toan/giao hang."],
        "main": [
            "Khach hang truy cap trang Lich su don hang.",
            "He thong tai danh sach don hang theo khach hang.",
            "Khach hang chon mot don hang de xem chi tiet.",
            "He thong hien thi ma don hang, san pham, tong tien, phi van chuyen, dia chi va trang thai.",
        ],
        "alt": [
            ("A1: Chua co don hang", ["He thong hien thi danh sach rong."]),
            ("A2: Don hang khong thuoc khach hang", ["He thong tu choi hien thi thong tin."]),
        ],
    },
    {
        "id": "UC10",
        "name": "Quan ly ho so va dia chi khach hang",
        "actor": "Khach hang",
        "description": "Cho phep khach hang xem thong tin ca nhan va quan ly dia chi giao hang.",
        "pre": ["Khach hang da dang nhap."],
        "post": ["Thong tin ho so/dia chi duoc cap nhat trong co so du lieu."],
        "main": [
            "Khach hang truy cap trang ho so.",
            "He thong hien thi thong tin tai khoan va danh sach dia chi.",
            "Khach hang them, sua hoac xoa dia chi.",
            "He thong kiem tra du lieu dia chi.",
            "He thong luu thay doi va cap nhat giao dien.",
        ],
        "alt": [
            ("A1: Thieu thong tin dia chi", ["He thong yeu cau nhap day du nguoi nhan, so dien thoai va dia chi."]),
            ("A2: Xoa dia chi dang duoc su dung", ["He thong co the tu choi xoa hoac yeu cau chon dia chi khac."]),
        ],
    },
    {
        "id": "UC11",
        "name": "Danh gia san pham/dich vu",
        "actor": "Khach hang",
        "description": "Cho phep khach hang tao, xem va xoa danh gia doi voi o to, phu kien hoac dich vu.",
        "pre": ["Khach hang da dang nhap khi tao hoac xoa danh gia.", "Doi tuong duoc danh gia ton tai."],
        "post": ["Danh gia duoc luu hoac xoa khoi he thong."],
        "main": [
            "Khach hang mo trang chi tiet san pham/dich vu.",
            "He thong hien thi danh sach danh gia hien co.",
            "Khach hang nhap diem danh gia va noi dung.",
            "Khach hang gui danh gia.",
            "He thong luu danh gia va cap nhat danh sach.",
        ],
        "alt": [
            ("A1: Chua dang nhap", ["He thong yeu cau dang nhap truoc khi gui danh gia."]),
            ("A2: Noi dung khong hop le", ["He thong hien thi thong bao loi validation."]),
        ],
    },
    {
        "id": "UC12",
        "name": "Gui khieu nai ho tro",
        "actor": "Khach hang",
        "description": "Cho phep khach hang tao khieu nai lien quan den don hang hoac trai nghiem su dung he thong.",
        "pre": ["Khach hang da dang nhap."],
        "post": ["Khieu nai duoc tao voi trang thai ban dau va cho nhan vien xu ly."],
        "main": [
            "Khach hang truy cap chuc nang ho tro/khieu nai.",
            "Khach hang nhap noi dung khieu nai va thong tin lien quan.",
            "Khach hang gui khieu nai.",
            "He thong kiem tra thong tin.",
            "He thong luu khieu nai voi trang thai moi.",
            "He thong thong bao gui khieu nai thanh cong.",
        ],
        "alt": [
            ("A1: Thieu noi dung", ["He thong yeu cau nhap noi dung khieu nai."]),
            ("A2: Loi xu ly", ["He thong thong bao khong the tao khieu nai va cho phep thu lai."]),
        ],
    },
    {
        "id": "UC13",
        "name": "Quan ly o to",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep nhan vien quan ly danh sach o to: them moi, cap nhat, doi trang thai, xoa va tai media.",
        "pre": ["Nguoi dung da dang nhap voi vai tro Admin hoac Nhan vien."],
        "post": ["Du lieu o to va media lien quan duoc cap nhat trong he thong."],
        "main": [
            "Nhan vien truy cap man hinh quan ly o to.",
            "He thong hien thi danh sach o to.",
            "Nhan vien them moi hoac chon mot o to de chinh sua.",
            "Nhan vien nhap thong tin xe, gia, hang xe, trang thai va hinh anh.",
            "He thong kiem tra va luu thong tin.",
            "He thong cap nhat danh sach o to.",
        ],
        "alt": [
            ("A1: Thieu thong tin bat buoc", ["He thong hien thi loi validation."]),
            ("A2: Khong co quyen xoa", ["He thong chi cho Admin thuc hien xoa o to."]),
            ("A3: Loi upload media", ["He thong thong bao loi va cho phep tai lai."]),
        ],
    },
    {
        "id": "UC14",
        "name": "Quan ly phu kien",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep quan ly phu kien: them, sua, xoa, tim kiem, phan loai va cap nhat so luong/trong luong.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien."],
        "post": ["Thong tin phu kien duoc cap nhat."],
        "main": [
            "Nhan vien truy cap man hinh quan ly phu kien.",
            "He thong hien thi danh sach phu kien.",
            "Nhan vien them moi hoac chinh sua phu kien.",
            "Nhan vien nhap ten, loai, hang san xuat, gia, so luong, trong luong va mo ta.",
            "He thong luu thong tin va cap nhat danh sach.",
        ],
        "alt": [
            ("A1: Thieu trong luong", ["He thong yeu cau nhap trong luong de tinh phi van chuyen GHN."]),
            ("A2: Khong co quyen xoa", ["He thong chi cho Admin xoa phu kien."]),
        ],
    },
    {
        "id": "UC15",
        "name": "Quan ly dich vu",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep quan ly dich vu: them moi, cap nhat, xoa va tim kiem dich vu.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien."],
        "post": ["Thong tin dich vu duoc cap nhat."],
        "main": [
            "Nhan vien truy cap man hinh quan ly dich vu.",
            "He thong hien thi danh sach dich vu.",
            "Nhan vien them moi hoac sua dich vu.",
            "Nhan vien nhap thong tin ten dich vu, gia, mo ta va trang thai.",
            "He thong luu thong tin va cap nhat danh sach.",
        ],
        "alt": [
            ("A1: Thieu thong tin", ["He thong hien thi thong bao yeu cau nhap day du thong tin."]),
            ("A2: Khong co quyen xoa", ["He thong chi cho Admin xoa dich vu."]),
        ],
    },
    {
        "id": "UC16",
        "name": "Quan ly don hang",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep nhan vien xem, tim kiem, cap nhat trang thai, gan nhan vien phu trach va xu ly van chuyen don hang.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien.", "Don hang ton tai trong he thong."],
        "post": ["Trang thai don hang duoc cap nhat.", "Don GHN duoc tao khi don hang duoc xac nhan neu cau hinh hop le."],
        "main": [
            "Nhan vien truy cap man hinh quan ly don hang.",
            "He thong hien thi danh sach don hang.",
            "Nhan vien xem chi tiet don hang.",
            "Nhan vien cap nhat trang thai don hang hoac gan nhan vien phu trach.",
            "Neu chuyen sang Da xac nhan, he thong tao don van chuyen GHN.",
            "He thong luu ma van don GHN va cap nhat trang thai.",
        ],
        "alt": [
            ("A1: Loi tao don GHN", ["He thong thong bao loi va giu don hang de nhan vien xu ly lai."]),
            ("A2: Trang thai khong hop le", ["He thong tu choi cap nhat trang thai."]),
        ],
    },
    {
        "id": "UC17",
        "name": "Quan ly lich hen",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep nhan vien xem lich lai thu, lich dich vu va cap nhat trang thai lich hen.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien.", "He thong co lich hen do khach hang tao."],
        "post": ["Trang thai lich hen duoc cap nhat."],
        "main": [
            "Nhan vien truy cap man hinh lich lai thu hoac lich dich vu.",
            "He thong hien thi danh sach lich hen theo loai.",
            "Nhan vien xem thong tin khach hang, san pham/dich vu, chi nhanh, ngay gio hen.",
            "Nhan vien cap nhat trang thai lich hen.",
            "He thong luu thay doi va cap nhat danh sach.",
        ],
        "alt": [
            ("A1: Lich hen khong ton tai", ["He thong hien thi thong bao khong tim thay."]),
            ("A2: Trang thai khong hop le", ["He thong tu choi cap nhat."]),
        ],
    },
    {
        "id": "UC18",
        "name": "Quan ly khach hang",
        "actor": "Quan tri vien",
        "description": "Cho phep quan tri vien xem danh sach khach hang, tim kiem va khoa/mo khoa tai khoan.",
        "pre": ["Nguoi dung da dang nhap voi quyen Admin."],
        "post": ["Thong tin trang thai tai khoan khach hang duoc cap nhat."],
        "main": [
            "Quan tri vien truy cap man hinh khach hang.",
            "He thong hien thi danh sach khach hang.",
            "Quan tri vien tim kiem khach hang theo tu khoa.",
            "Quan tri vien cap nhat trang thai tai khoan.",
            "He thong luu thay doi va cap nhat danh sach.",
        ],
        "alt": [
            ("A1: Khong tim thay khach hang", ["He thong hien thi ket qua rong."]),
            ("A2: Khong co quyen", ["He thong tu choi truy cap neu nguoi dung khong phai Admin."]),
        ],
    },
    {
        "id": "UC19",
        "name": "Quan ly nhan vien",
        "actor": "Quan tri vien",
        "description": "Cho phep quan tri vien them, xem, sua, xoa, tim kiem nhan vien va loc theo chuc vu.",
        "pre": ["Quan tri vien da dang nhap."],
        "post": ["Du lieu nhan vien duoc cap nhat trong he thong."],
        "main": [
            "Quan tri vien truy cap chuc nang quan ly nhan vien.",
            "He thong hien thi danh sach nhan vien.",
            "Quan tri vien them moi hoac chinh sua thong tin nhan vien.",
            "He thong kiem tra du lieu.",
            "He thong luu thong tin nhan vien.",
        ],
        "alt": [
            ("A1: Email/tai khoan nhan vien da ton tai", ["He thong thong bao loi trung du lieu."]),
            ("A2: Thieu thong tin bat buoc", ["He thong yeu cau nhap day du thong tin."]),
        ],
    },
    {
        "id": "UC20",
        "name": "Quan ly kho hang va ton kho",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep xem danh sach kho, chi tiet kho va cap nhat ton kho cho o to/phu kien tai tung kho.",
        "pre": ["Nguoi dung co quyen quan tri hoac nhan vien.", "Kho hang da duoc cau hinh."],
        "post": ["Du lieu ton kho duoc cap nhat."],
        "main": [
            "Nhan vien truy cap man hinh kho hang hoac ton kho.",
            "He thong hien thi danh sach kho va so lieu ton kho.",
            "Nhan vien chon san pham can cap nhat ton kho.",
            "Nhan vien nhap so luong ton tai kho.",
            "He thong luu thong tin ton kho.",
        ],
        "alt": [
            ("A1: So luong am hoac khong hop le", ["He thong tu choi luu va hien thi thong bao loi."]),
            ("A2: Kho khong ton tai", ["He thong thong bao khong tim thay kho."]),
        ],
    },
    {
        "id": "UC21",
        "name": "Quan ly danh gia",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep nhan vien xem danh sach danh gia va xoa cac danh gia khong phu hop.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien."],
        "post": ["Danh gia vi pham duoc xoa khoi he thong neu can."],
        "main": [
            "Nhan vien truy cap man hinh danh gia.",
            "He thong hien thi danh sach danh gia.",
            "Nhan vien xem noi dung, diem so va doi tuong duoc danh gia.",
            "Nhan vien xoa danh gia neu khong phu hop.",
            "He thong cap nhat danh sach danh gia.",
        ],
        "alt": [
            ("A1: Danh gia khong ton tai", ["He thong hien thi thong bao khong tim thay."]),
            ("A2: Khong co quyen xoa", ["He thong tu choi thao tac."]),
        ],
    },
    {
        "id": "UC22",
        "name": "Xu ly khieu nai",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep nhan vien xem khieu nai, phan hoi khach hang va cap nhat trang thai xu ly.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien.", "Khieu nai ton tai trong he thong."],
        "post": ["Khieu nai duoc phan hoi va cap nhat trang thai."],
        "main": [
            "Nhan vien truy cap man hinh khieu nai.",
            "He thong hien thi danh sach khieu nai.",
            "Nhan vien mo chi tiet khieu nai.",
            "Nhan vien nhap noi dung phan hoi.",
            "Nhan vien cap nhat trang thai khieu nai.",
            "He thong luu phan hoi va trang thai moi.",
        ],
        "alt": [
            ("A1: Thieu noi dung phan hoi", ["He thong yeu cau nhap phan hoi."]),
            ("A2: Trang thai khong hop le", ["He thong tu choi cap nhat."]),
        ],
    },
    {
        "id": "UC23",
        "name": "Quan ly media san pham",
        "actor": "Quan tri vien, Nhan vien",
        "description": "Cho phep tai len, xem va xoa hinh anh/video gan voi o to, phu kien hoac dich vu.",
        "pre": ["Nguoi dung co quyen Admin hoac Nhan vien khi tai/xoa media.", "Doi tuong can gan media ton tai."],
        "post": ["Media duoc luu, hien thi cong khai hoac bi xoa khoi doi tuong."],
        "main": [
            "Nhan vien mo man hinh chi tiet san pham/dich vu.",
            "Nhan vien chon tep hinh anh hoac video.",
            "He thong tai media len va gan voi doi tuong.",
            "He thong hien thi media moi trong danh sach.",
            "Nhan vien co the xoa media neu can.",
        ],
        "alt": [
            ("A1: Tep khong hop le", ["He thong tu choi tai len va hien thi thong bao loi."]),
            ("A2: Loi dich vu luu tru", ["He thong thong bao khong the tai media."]),
        ],
    },
    {
        "id": "UC24",
        "name": "Tra cuu va tinh phi van chuyen GHN",
        "actor": "Khach hang, He thong, Nhan vien",
        "description": "Cho phep he thong tra cuu tinh/thanh, quan/huyen, phuong/xa va tinh phi van chuyen qua GHN trong luong thanh toan.",
        "pre": ["Cau hinh GHN hop le hoac co co che fallback demo.", "Nguoi dung dang o man hinh can thong tin giao hang."],
        "post": ["Dia chi GHN va phi van chuyen duoc tra ve cho quy trinh dat hang."],
        "main": [
            "Nguoi dung chon tinh/thanh trong form dia chi.",
            "He thong tai danh sach quan/huyen tu GHN.",
            "Nguoi dung chon quan/huyen.",
            "He thong tai danh sach phuong/xa.",
            "Nguoi dung chon phuong/xa.",
            "He thong tinh phi van chuyen dua tren kho xuat hang, dia chi nhan va trong luong.",
            "He thong hien thi phi van chuyen tren man hinh thanh toan.",
        ],
        "alt": [
            ("A1: GHN khong phan hoi", ["He thong hien thi loi hoac su dung phi mac dinh theo cau hinh demo."]),
            ("A2: Dia chi khong hop le", ["He thong yeu cau nguoi dung chon lai dia chi."]),
        ],
    },
    {
        "id": "UC25",
        "name": "Tiep nhan webhook GHN",
        "actor": "He thong GHN",
        "description": "Cho phep GHN gui cap nhat trang thai van don ve he thong de dong bo trang thai giao hang.",
        "pre": ["Don hang da co ma van don GHN.", "Webhook GHN da duoc cau hinh."],
        "post": ["Trang thai van chuyen/don hang duoc cap nhat theo du lieu GHN."],
        "main": [
            "GHN gui webhook cap nhat trang thai van don.",
            "He thong tiep nhan payload webhook.",
            "He thong xac dinh don hang tu ma van don.",
            "He thong cap nhat trang thai giao hang tuong ung.",
            "He thong tra ve phan hoi thanh cong cho GHN.",
        ],
        "alt": [
            ("A1: Khong tim thay ma van don", ["He thong ghi nhan loi va khong cap nhat don hang."]),
            ("A2: Payload khong hop le", ["He thong tu choi xu ly webhook."]),
        ],
    },
]


def set_default_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ĐẶC TẢ USE CASE HỆ THỐNG CARSHOP")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dự án: Website thương mại điện tử ô tô")
    run.italic = True
    run.font.size = Pt(13)

    document.add_paragraph()


def add_overview(document: Document) -> None:
    document.add_heading("1. Tổng quan", level=1)
    document.add_paragraph(
        "Tài liệu này đặc tả các use case chính được suy ra từ cấu trúc frontend, "
        "các route React, controller backend Spring Boot và tài liệu API hiện có của dự án CarShop. "
        "Các use case được gom theo nghiệp vụ người dùng thay vì liệt kê từng endpoint riêng lẻ."
    )
    document.add_paragraph("Phạm vi tác nhân:")
    for actor in [
        "Khách vãng lai: xem, tìm kiếm sản phẩm/dịch vụ.",
        "Khách hàng: đăng ký, đăng nhập, quản lý giỏ hàng, đặt hàng, đặt lịch, đánh giá, khiếu nại.",
        "Nhân viên: quản lý sản phẩm, đơn hàng, lịch hẹn, đánh giá, khiếu nại, kho/tồn kho theo phân quyền.",
        "Quản trị viên: toàn quyền quản trị, bao gồm quản lý khách hàng và nhân viên.",
        "Hệ thống GHN/VNPay: tích hợp vận chuyển và thanh toán.",
    ]:
        document.add_paragraph(actor, style="List Bullet")


def add_summary_table(document: Document) -> None:
    document.add_heading("2. Danh sách use case", level=1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Mã", "Tên use case", "Tác nhân chính", "Nhóm chức năng"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header

    for uc in use_cases:
        row = table.add_row().cells
        row[0].text = uc["id"]
        row[1].text = uc["name"]
        row[2].text = uc["actor"]
        if int(uc["id"][2:]) <= 12:
            group = "Khách hàng"
        elif int(uc["id"][2:]) <= 23:
            group = "Quản trị/Nhân viên"
        else:
            group = "Tích hợp hệ thống"
        row[3].text = group


def add_list(document: Document, items, numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        document.add_paragraph(item, style=style)


def add_use_case(document: Document, uc: dict) -> None:
    document.add_heading(f"{uc['id']} - {uc['name']}", level=2)

    table = document.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Tác nhân", uc["actor"]),
        ("Mô tả", uc["description"]),
        ("Tiền điều kiện", "\n".join(uc["pre"])),
        ("Hậu điều kiện", "\n".join(uc["post"])),
        ("Ghi chú liên quan", "Đặc tả dựa trên route frontend và API/backend hiện có trong dự án."),
    ]
    for i, (key, value) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value

    document.add_heading("Luồng sự kiện chính", level=3)
    add_list(document, uc["main"], numbered=True)

    document.add_heading("Luồng sự kiện phụ", level=3)
    for title, steps in uc["alt"]:
        p = document.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        add_list(document, steps, numbered=True)


def build_document() -> None:
    document = Document()
    set_default_styles(document)
    add_title(document)
    add_overview(document)
    add_summary_table(document)
    document.add_page_break()
    document.add_heading("3. Đặc tả chi tiết use case", level=1)
    for idx, uc in enumerate(use_cases):
        add_use_case(document, uc)
        if idx != len(use_cases) - 1:
            document.add_page_break()
    document.save(OUTPUT)


if __name__ == "__main__":
    vietnamese_names = {
        "UC01": "Đăng ký tài khoản",
        "UC02": "Đăng nhập hệ thống",
        "UC03": "Xem và tìm kiếm ô tô",
        "UC04": "Xem và tìm kiếm phụ kiện",
        "UC05": "Xem dịch vụ và đặt lịch dịch vụ",
        "UC06": "Đăng ký lái thử ô tô",
        "UC07": "Quản lý giỏ hàng",
        "UC08": "Đặt hàng và thanh toán",
        "UC09": "Xem lịch sử và chi tiết đơn hàng",
        "UC10": "Quản lý hồ sơ và địa chỉ khách hàng",
        "UC11": "Đánh giá sản phẩm/dịch vụ",
        "UC12": "Gửi khiếu nại hỗ trợ",
        "UC13": "Quản lý ô tô",
        "UC14": "Quản lý phụ kiện",
        "UC15": "Quản lý dịch vụ",
        "UC16": "Quản lý đơn hàng",
        "UC17": "Quản lý lịch hẹn",
        "UC18": "Quản lý khách hàng",
        "UC19": "Quản lý nhân viên",
        "UC20": "Quản lý kho hàng và tồn kho",
        "UC21": "Quản lý đánh giá",
        "UC22": "Xử lý khiếu nại",
        "UC23": "Quản lý media sản phẩm",
        "UC24": "Tra cứu và tính phí vận chuyển GHN",
        "UC25": "Tiếp nhận webhook GHN",
    }
    for uc in use_cases:
        if uc["id"] in vietnamese_names:
            uc["name"] = vietnamese_names[uc["id"]]

    # Replace the first use case with the fully accented text supplied by the user.
    use_cases[0].update({
        "actor": "Người dùng (Khách hàng)",
        "description": "Use Case này cho phép người dùng nhập thông tin cá nhân cần thiết để tạo một tài khoản mới và đăng ký sử dụng hệ thống.",
        "pre": ["Người dùng chưa có tài khoản trong hệ thống."],
        "post": [
            "Tài khoản mới được tạo trong hệ thống.",
            "Thông tin người dùng được lưu vào cơ sở dữ liệu.",
        ],
        "main": [
            "Người dùng truy cập trang đăng ký.",
            "Hệ thống hiển thị form đăng ký tài khoản.",
            "Người dùng nhập các thông tin cần thiết: họ tên, email, mật khẩu, xác nhận mật khẩu, số điện thoại.",
            "Người dùng nhấn nút Đăng ký.",
            "Hệ thống kiểm tra tính hợp lệ của thông tin đã nhập.",
            "Nếu thông tin hợp lệ, hệ thống tạo tài khoản mới trong cơ sở dữ liệu.",
            "Hệ thống thông báo đăng ký thành công.",
            "Hệ thống chuyển người dùng đến trang đăng nhập.",
        ],
        "alt": [
            ("A1: Tên đăng nhập hoặc email đã tồn tại", [
                "Hệ thống kiểm tra thông tin đăng ký.",
                "Phát hiện tên đăng nhập hoặc email đã được sử dụng.",
                "Hệ thống hiển thị thông báo: Tên đăng nhập hoặc email đã tồn tại.",
                "Người dùng nhập lại thông tin khác.",
            ]),
            ("A2: Chưa nhập đầy đủ thông tin", [
                "Người dùng nhấn Đăng ký nhưng để trống một hoặc nhiều trường.",
                "Hệ thống hiển thị thông báo: Vui lòng nhập đầy đủ thông tin.",
            ]),
            ("A3: Mật khẩu và xác nhận mật khẩu không khớp", [
                "Hệ thống kiểm tra thông tin mật khẩu.",
                "Phát hiện mật khẩu và xác nhận mật khẩu không giống nhau.",
                "Hệ thống hiển thị thông báo: Mật khẩu xác nhận không khớp.",
            ]),
        ],
    })

    common_actor_replacements = {
        "Khach hang": "Khách hàng",
        "Khach vang lai": "Khách vãng lai",
        "Nhan vien": "Nhân viên",
        "Quan tri vien": "Quản trị viên",
        "He thong": "Hệ thống",
        "He thong GHN": "Hệ thống GHN",
    }
    for uc in use_cases[1:]:
        for src, dst in common_actor_replacements.items():
            uc["actor"] = uc["actor"].replace(src, dst)

    accent_replacements = [
        ("Use Case nay", "Use Case này"),
        ("Cho phep", "Cho phép"),
        ("cho phep", "cho phép"),
        ("nguoi dung", "người dùng"),
        ("Nguoi dung", "Người dùng"),
        ("khach hang", "khách hàng"),
        ("Khach hang", "Khách hàng"),
        ("khach vang lai", "khách vãng lai"),
        ("Khach vang lai", "Khách vãng lai"),
        ("nhan vien", "nhân viên"),
        ("Nhan vien", "Nhân viên"),
        ("quan tri vien", "quản trị viên"),
        ("Quan tri vien", "Quản trị viên"),
        ("he thong", "hệ thống"),
        ("He thong", "Hệ thống"),
        ("dang nhap", "đăng nhập"),
        ("Dang nhap", "Đăng nhập"),
        ("dang ky", "đăng ký"),
        ("Dang ky", "Đăng ký"),
        ("tai khoan", "tài khoản"),
        ("Tai khoan", "Tài khoản"),
        ("thong tin", "thông tin"),
        ("Thong tin", "Thông tin"),
        ("co so du lieu", "cơ sở dữ liệu"),
        ("trang chu", "trang chủ"),
        ("trang quan tri", "trang quản trị"),
        ("mat khau", "mật khẩu"),
        ("Mat khau", "Mật khẩu"),
        ("email", "email"),
        ("so dien thoai", "số điện thoại"),
        ("So dien thoai", "Số điện thoại"),
        ("ho ten", "họ tên"),
        ("Ho ten", "Họ tên"),
        ("hop le", "hợp lệ"),
        ("khong hop le", "không hợp lệ"),
        ("khong", "không"),
        ("Khong", "Không"),
        ("ton tai", "tồn tại"),
        ("Ton tai", "Tồn tại"),
        ("duoc", "được"),
        ("Duoc", "Được"),
        ("luu", "lưu"),
        ("Luu", "Lưu"),
        ("cap nhat", "cập nhật"),
        ("Cap nhat", "Cập nhật"),
        ("hien thi", "hiển thị"),
        ("Hien thi", "Hiển thị"),
        ("thong bao", "thông báo"),
        ("Thong bao", "Thông báo"),
        ("yeu cau", "yêu cầu"),
        ("Yeu cau", "Yêu cầu"),
        ("truy cap", "truy cập"),
        ("Truy cap", "Truy cập"),
        ("nhap", "nhập"),
        ("Nhap", "Nhập"),
        ("nhan", "nhấn"),
        ("Nhan", "Nhấn"),
        ("kiem tra", "kiểm tra"),
        ("Kiem tra", "Kiểm tra"),
        ("xac thuc", "xác thực"),
        ("Xac thuc", "Xác thực"),
        ("tao", "tạo"),
        ("Tao", "Tạo"),
        ("chuyen", "chuyển"),
        ("Chuyen", "Chuyển"),
        ("nut", "nút"),
        ("Nut", "Nút"),
        ("Neu", "Nếu"),
        ("neu", "nếu"),
        ("den", "đến"),
        ("Den", "Đến"),
        ("hoac", "hoặc"),
        ("Hoac", "Hoặc"),
        (" va ", " và "),
        ("co ", "có "),
        (" co", " có"),
        ("chon", "chọn"),
        ("Chon", "Chọn"),
        ("ket qua", "kết quả"),
        ("Ket qua", "Kết quả"),
        ("phu hop", "phù hợp"),
        ("Phu hop", "Phù hợp"),
        ("tim thay", "tìm thấy"),
        ("Tim thay", "Tìm thấy"),
        ("rong", "rỗng"),
        ("Rong", "Rỗng"),
        ("het hang", "hết hàng"),
        ("Het hang", "Hết hàng"),
        ("van ", "vẫn "),
        ("Van ", "Vẫn "),
        ("nhung", "nhưng"),
        ("Nhung", "Nhưng"),
        ("dat", "đặt"),
        ("Dat", "Đặt"),
        (" qua ", " quá "),
        ("thu lai", "thử lại"),
        ("Thu lai", "Thử lại"),
        ("tu choi", "từ chối"),
        ("Tu choi", "Từ chối"),
        ("bi khoa", "bị khóa"),
        ("Bi khoa", "Bị khóa"),
        ("kha dung", "khả dụng"),
        ("Kha dung", "Khả dụng"),
        ("phan trang", "phân trang"),
        ("Phan trang", "Phân trang"),
        ("gui", "gửi"),
        ("Gui", "Gửi"),
        ("loc", "lọc"),
        ("Loc", "Lọc"),
        ("mot ", "một "),
        ("Mot ", "Một "),
        (" de ", " để "),
        ("tinh trang", "tình trạng"),
        ("Tinh trang", "Tình trạng"),
        ("tai ", "tải "),
        ("Tai ", "Tải "),
        ("xem", "xem"),
        ("Xem", "Xem"),
        ("tim kiem", "tìm kiếm"),
        ("Tim kiem", "Tìm kiếm"),
        ("danh sach", "danh sách"),
        ("Danh sach", "Danh sách"),
        ("chi tiet", "chi tiết"),
        ("Chi tiet", "Chi tiết"),
        ("san pham", "sản phẩm"),
        ("San pham", "Sản phẩm"),
        ("dich vu", "dịch vụ"),
        ("Dich vu", "Dịch vụ"),
        ("phu kien", "phụ kiện"),
        ("Phu kien", "Phụ kiện"),
        ("o to", "ô tô"),
        ("O to", "Ô tô"),
        ("gio hang", "giỏ hàng"),
        ("Gio hang", "Giỏ hàng"),
        ("don hang", "đơn hàng"),
        ("Don hang", "Đơn hàng"),
        ("dia chi", "địa chỉ"),
        ("Dia chi", "Địa chỉ"),
        ("thanh toan", "thanh toán"),
        ("Thanh toan", "Thanh toán"),
        ("van chuyen", "vận chuyển"),
        ("Van chuyen", "Vận chuyển"),
        ("phi", "phí"),
        ("Phi", "Phí"),
        ("lich hen", "lịch hẹn"),
        ("Lich hen", "Lịch hẹn"),
        ("lai thu", "lái thử"),
        ("Lai thu", "Lái thử"),
        ("danh gia", "đánh giá"),
        ("Danh gia", "Đánh giá"),
        ("khieu nai", "khiếu nại"),
        ("Khieu nai", "Khiếu nại"),
        ("trang thai", "trạng thái"),
        ("Trang thai", "Trạng thái"),
        ("vai tro", "vai trò"),
        ("Vai tro", "Vai trò"),
        ("quyen", "quyền"),
        ("Quyen", "Quyền"),
        ("kho", "kho"),
        ("ton kho", "tồn kho"),
        ("Ton kho", "Tồn kho"),
        ("ngay", "ngày"),
        ("Ngay", "Ngày"),
        ("gio", "giờ"),
        ("Gio", "Giờ"),
        ("ghi chu", "ghi chú"),
        ("Ghi chu", "Ghi chú"),
        ("noi dung", "nội dung"),
        ("Noi dung", "Nội dung"),
        ("phan hoi", "phản hồi"),
        ("Phan hoi", "Phản hồi"),
        ("loi", "lỗi"),
        ("Loi", "Lỗi"),
        ("bo loc", "bộ lọc"),
        ("Bo loc", "Bộ lọc"),
        ("tu khoa", "từ khóa"),
        ("Tu khoa", "Từ khóa"),
        ("gia", "giá"),
        ("Gia", "Giá"),
        ("so luong", "số lượng"),
        ("So luong", "Số lượng"),
        ("hinh anh", "hình ảnh"),
        ("Hinh anh", "Hình ảnh"),
        ("video", "video"),
        ("ma", "mã"),
        ("Ma", "Mã"),
        ("du lieu", "dữ liệu"),
        ("Du lieu", "Dữ liệu"),
        ("cau hinh", "cấu hình"),
        ("Cau hinh", "Cấu hình"),
        ("xu ly", "xử lý"),
        ("Xu ly", "Xử lý"),
    ]

    def accent_text(value):
        if isinstance(value, str):
            for src, dst in accent_replacements:
                value = value.replace(src, dst)
            value = value.replace("emãil", "email").replace("Emãil", "Email")
            return value
        if isinstance(value, list):
            return [accent_text(item) for item in value]
        if isinstance(value, tuple):
            return tuple(accent_text(item) for item in value)
        return value

    for uc in use_cases[1:]:
        for key in ["actor", "description", "pre", "post", "main", "alt"]:
            uc[key] = accent_text(uc[key])

    build_document()
    print(OUTPUT)
